"""
document_processor.py - Template Style Extractor
=================================================
Reads the user's uploaded .docx CV and cover letter templates and extracts:
  - Text content (for Claude to understand the candidate's history/voice)
  - Style fingerprint (fonts, sizes, spacing, colours, structure)

The style fingerprint is stored as JSON and passed to doc_generator.py
so every output document mirrors the template's look and feel.
"""

import json
import zipfile
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field, asdict

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class StyleFingerprint:
    """Captures the visual style of a document template."""

    # Fonts
    body_font: str = "Calibri"
    heading_font: str = "Calibri"
    body_size_pt: float = 11.0
    heading1_size_pt: float = 16.0
    heading2_size_pt: float = 13.0

    # Colours (hex strings, e.g. "2E74B5")
    heading_colour: Optional[str] = None
    accent_colour: Optional[str] = None

    # Spacing
    paragraph_space_before_pt: float = 0.0
    paragraph_space_after_pt: float = 6.0
    line_spacing: Optional[float] = None  # None = single

    # Margins (in cm)
    margin_top_cm: float = 2.54
    margin_bottom_cm: float = 2.54
    margin_left_cm: float = 2.54
    margin_right_cm: float = 2.54

    # Structure detected in the document
    has_header_block: bool = True         # Name/contact at top
    has_horizontal_rule: bool = False     # Divider lines between sections
    section_headings: List[str] = field(default_factory=list)

    # Cover letter specific
    uses_date_line: bool = True
    uses_salutation: bool = True
    sign_off: str = "Yours sincerely"

    def to_json(self) -> str:
        return json.dumps(asdict(self), indent=2)

    @classmethod
    def from_json(cls, s: str) -> "StyleFingerprint":
        return cls(**json.loads(s))


class DocumentProcessor:
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )

    def extract_cv_template(self, docx_path: Path) -> tuple[str, StyleFingerprint]:
        """
        Returns (full_text_content, style_fingerprint) from the uploaded CV.
        full_text_content is passed to Claude as the base CV.
        style_fingerprint drives formatting of output documents.
        """
        doc = Document(str(docx_path))
        text = self._extract_text(doc)
        style = self._extract_style(doc, document_type="cv")
        return text, style

    def extract_cover_letter_template(self, docx_path: Path) -> tuple[str, StyleFingerprint]:
        """
        Returns (full_text_content, style_fingerprint) from the uploaded cover letter.
        """
        doc = Document(str(docx_path))
        text = self._extract_text(doc)
        style = self._extract_style(doc, document_type="cover_letter")
        # Detect sign-off
        style.sign_off = self._detect_sign_off(text)
        style.uses_salutation = self._detect_salutation(text)
        return text, style

    def _extract_text(self, doc) -> str:
        """Extract all paragraph text, preserving structure."""
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        return "\n".join(lines)

    def _extract_style(self, doc, document_type: str = "cv") -> StyleFingerprint:
        fp = StyleFingerprint()

        # ── Page margins ──────────────────────────────────────────────────────
        for section in doc.sections:
            fp.margin_top_cm = round(section.top_margin.cm, 2) if section.top_margin else 2.54
            fp.margin_bottom_cm = round(section.bottom_margin.cm, 2) if section.bottom_margin else 2.54
            fp.margin_left_cm = round(section.left_margin.cm, 2) if section.left_margin else 2.54
            fp.margin_right_cm = round(section.right_margin.cm, 2) if section.right_margin else 2.54
            break  # only need first section

        # ── Font and size from paragraphs ─────────────────────────────────────
        body_fonts = []
        heading_fonts = []
        body_sizes = []
        h1_sizes = []
        h2_sizes = []

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            is_heading1 = "heading 1" in style_name
            is_heading2 = "heading 2" in style_name
            is_body = not is_heading1 and not is_heading2

            for run in para.runs:
                if not run.text.strip():
                    continue
                font_name = run.font.name or (para.style.font.name if para.style else None)
                font_size = run.font.size or (para.style.font.size if para.style else None)
                size_pt = round(font_size.pt, 1) if font_size else None

                if is_heading1:
                    if font_name:
                        heading_fonts.append(font_name)
                    if size_pt:
                        h1_sizes.append(size_pt)
                    # Capture heading colour
                    if run.font.color and run.font.color.type is not None:
                        try:
                            rgb = run.font.color.rgb
                            if rgb:
                                fp.heading_colour = str(rgb)
                        except Exception:
                            pass
                elif is_heading2:
                    if font_name:
                        heading_fonts.append(font_name)
                    if size_pt:
                        h2_sizes.append(size_pt)
                else:
                    if font_name:
                        body_fonts.append(font_name)
                    if size_pt:
                        body_sizes.append(size_pt)

        fp.body_font = self._most_common(body_fonts) or "Calibri"
        fp.heading_font = self._most_common(heading_fonts) or fp.body_font
        fp.body_size_pt = self._median(body_sizes) or 11.0
        fp.heading1_size_pt = self._median(h1_sizes) or 16.0
        fp.heading2_size_pt = self._median(h2_sizes) or 13.0

        # ── Paragraph spacing ─────────────────────────────────────────────────
        space_afters = []
        for para in doc.paragraphs:
            if para.paragraph_format.space_after:
                space_afters.append(para.paragraph_format.space_after.pt)
        if space_afters:
            fp.paragraph_space_after_pt = round(self._median(space_afters), 1)

        # ── Section headings (for CV structure) ───────────────────────────────
        fp.section_headings = []
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            if ("heading" in style_name or para.text.isupper()) and para.text.strip():
                fp.section_headings.append(para.text.strip())

        # ── Horizontal rules ──────────────────────────────────────────────────
        # Detect if borders are used as section separators
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.element is not None:
                pPr = pf.element.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
                )
                if pPr is not None:
                    fp.has_horizontal_rule = True
                    break

        return fp

    @staticmethod
    def _detect_sign_off(text: str) -> str:
        for phrase in ["Yours sincerely", "Kind regards", "Best regards", "Yours faithfully", "Sincerely"]:
            if phrase.lower() in text.lower():
                return phrase
        return "Yours sincerely"

    @staticmethod
    def _detect_salutation(text: str) -> bool:
        return any(
            kw in text.lower() for kw in ["dear hiring", "dear recruiter", "dear sir", "dear madam", "to whom"]
        )

    @staticmethod
    def _most_common(lst: list):
        if not lst:
            return None
        return max(set(lst), key=lst.count)

    @staticmethod
    def _median(lst: list):
        if not lst:
            return None
        s = sorted(lst)
        n = len(s)
        return s[n // 2] if n % 2 else (s[n // 2 - 1] + s[n // 2]) / 2
