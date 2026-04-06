"""
doc_generator.py - Styled .docx Generator
==========================================
Takes Claude's tailored Markdown text and renders it as a properly
formatted .docx document matching the user's uploaded template style.

Requires: python-docx
"""

import re
import sys
from pathlib import Path
from typing import Optional
from document_processor import StyleFingerprint

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocGenerator:
    """Converts Markdown text → styled .docx."""

    def __init__(self, style: Optional[StyleFingerprint] = None):
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        self.style = style or StyleFingerprint()

    def generate_cv(self, markdown_text: str, output_path: Path, ats_mode: bool = True) -> Path:
        """
        Generate a styled CV .docx from Markdown.
        ats_mode=True (default): ATS-safe output — no tables, plain bullets,
        contact details in body, standard headings only.
        """
        # Strip the AI-generated header comments
        markdown_text = re.sub(r"<!--.*?-->", "", markdown_text, flags=re.DOTALL).strip()
        # Strip the review warning blockquote
        markdown_text = re.sub(r"^>.*\n?", "", markdown_text, flags=re.MULTILINE).strip()

        doc = self._new_document()
        self._apply_margins(doc)
        self._render_markdown(doc, markdown_text, document_type="cv", ats_mode=ats_mode)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def generate_cover_letter(self, markdown_text: str, output_path: Path,
                               candidate_name: str = "", job_title: str = "",
                               employer: str = "", ats_mode: bool = True) -> Path:
        """Generate a styled cover letter .docx from Markdown. Returns path."""
        doc = self._new_document()
        self._apply_margins(doc)
        self._render_markdown(doc, markdown_text, document_type="cover_letter", ats_mode=ats_mode)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _new_document(self) -> "Document":
        doc = Document()
        # Remove default empty paragraph
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        return doc

    def _apply_margins(self, doc):
        s = self.style
        for section in doc.sections:
            section.top_margin = Cm(s.margin_top_cm)
            section.bottom_margin = Cm(s.margin_bottom_cm)
            section.left_margin = Cm(s.margin_left_cm)
            section.right_margin = Cm(s.margin_right_cm)

    def _render_markdown(self, doc, text: str, document_type: str = "cv", ats_mode: bool = True):
        """
        Parse simple Markdown and add styled paragraphs to doc.
        ats_mode=True enforces ATS-safe formatting:
          - No tables or columns
          - Plain hyphen bullets only
          - No decorative horizontal rules
          - Contact details rendered as plain body text
        """
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Right-aligned line (>>> prefix) — sender address in cover letters
            if line.startswith(">>> ") or line.strip() == ">>>":
                text = line[4:].strip() if line.startswith(">>> ") else ""
                self._add_right_aligned(doc, text)

            # H1 — candidate name or document title
            elif line.startswith("# "):
                self._add_heading(doc, line[2:].strip(), level=1)

            # H2 — section headers (EDUCATION, EXPERIENCE etc.)
            elif line.startswith("## "):
                self._add_section_heading(doc, line[3:].strip(), ats_mode=ats_mode)

            # H3 — job title / role title
            elif line.startswith("### "):
                self._add_heading(doc, line[4:].strip(), level=3)

            # Horizontal rule — skip in ATS mode (decorative only)
            elif line.strip() in ("---", "***", "___"):
                if not ats_mode and self.style.has_horizontal_rule:
                    self._add_rule(doc)

            # Bullet point — always plain hyphen in ATS mode
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                bullet_content = line.strip()[2:]
                self._add_bullet(doc, bullet_content, ats_mode=ats_mode)

            # Blockquote (strip — used for AI warnings)
            elif line.strip().startswith(">"):
                pass  # skip

            # Contact line — in ATS mode render as plain body text, not styled
            elif line.strip().startswith("**") and "|" in line:
                if ats_mode:
                    plain = re.sub(r"\*\*", "", line.strip())
                    self._add_body(doc, plain)
                else:
                    self._add_contact_line(doc, line.strip())

            # Empty line → small spacer
            elif line.strip() == "":
                self._add_spacer(doc)

            # Regular paragraph
            else:
                self._add_body(doc, line.strip())

            i += 1

    def _add_heading(self, doc, text: str, level: int = 1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.style.heading_font
        run.bold = True

        if level == 1:
            run.font.size = Pt(self.style.heading1_size_pt)
            if self.style.heading_colour:
                try:
                    hex_col = self.style.heading_colour.lstrip("#")
                    r, g, b = int(hex_col[0:2], 16), int(hex_col[2:4], 16), int(hex_col[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except Exception:
                    pass
        elif level == 3:
            run.font.size = Pt(self.style.body_size_pt + 1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

    def _add_section_heading(self, doc, text: str, ats_mode: bool = True):
        """H2 — styled as a section divider matching the template."""
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if text.isupper() or all(
            c.isupper() or not c.isalpha() for c in text
        ) else text)
        run.font.name = self.style.heading_font
        run.font.size = Pt(self.style.heading2_size_pt)
        run.bold = True

        if self.style.heading_colour:
            try:
                hex_col = self.style.heading_colour.lstrip("#")
                r, g, b = int(hex_col[0:2], 16), int(hex_col[2:4], 16), int(hex_col[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass

        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

        # Add bottom border if the template uses rules
        if self.style.has_horizontal_rule:
            self._add_para_bottom_border(p)

    def _add_bullet(self, doc, text: str, ats_mode: bool = True):
        if ats_mode:
            # ATS-safe: plain paragraph starting with hyphen-space
            # Many ATS systems parse "List Bullet" style unreliably;
            # a plain paragraph with a hyphen prefix is most universally parsed
            p = doc.add_paragraph()
            run = p.add_run("- " + re.sub(r"\*\*(.+?)\*\*", r"\1", text))  # strip bold in ATS mode
            run.font.name = self.style.body_font
            run.font.size = Pt(self.style.body_size_pt)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)  # add this line

        else:
            p = doc.add_paragraph(style="List Bullet")
            self._add_inline_markdown(p, text)
            run_base = p.runs[0] if p.runs else p.add_run()
            run_base.font.name = self.style.body_font
            run_base.font.size = Pt(self.style.body_size_pt)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)  # add this line

    def _add_body(self, doc, text: str):
        if not text:
            return
        p = doc.add_paragraph()
        self._add_inline_markdown(p, text)
        p.paragraph_format.space_after = Pt(self.style.paragraph_space_after_pt)

    def _add_right_aligned(self, doc, text: str):
        """Render a line in the right-hand column — sender address in cover letters.
        All lines start at the same left edge (half the text width) so the block
        sits consistently on the right side of the page."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Indent to 62% of the usable text area so the block sits further right
        text_width_cm = 21.0 - self.style.margin_left_cm - self.style.margin_right_cm
        p.paragraph_format.left_indent = Cm(text_width_cm * 0.62)
        if text:
            self._add_inline_markdown(p, text)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    def _add_contact_line(self, doc, text: str):
        """Render the contact/header line — bold labels, normal values."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_inline_markdown(p, text)
        p.paragraph_format.space_after = Pt(2)

    def _add_spacer(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(4)

    def _add_rule(self, doc):
        p = doc.add_paragraph()
        self._add_para_bottom_border(p)
        p.paragraph_format.space_after = Pt(4)

    def _add_inline_markdown(self, para, text: str):
        """Parse **bold**, *italic*, and plain text within a paragraph."""
        # Pattern matches **bold**, *italic*, or plain text
        pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)")
        for match in pattern.finditer(text):
            chunk = match.group()
            run = para.add_run()
            run.font.name = self.style.body_font
            run.font.size = Pt(self.style.body_size_pt)
            if chunk.startswith("**") and chunk.endswith("**"):
                run.bold = True
                run.text = chunk[2:-2]
            elif chunk.startswith("*") and chunk.endswith("*"):
                run.italic = True
                run.text = chunk[1:-1]
            else:
                run.text = chunk

    @staticmethod
    def _add_para_bottom_border(para):
        """Add a bottom border line to a paragraph (acts as a horizontal rule)."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

